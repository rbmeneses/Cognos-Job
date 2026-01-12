# Arquivo/Snippet [CognosJob.py]:
import streamlit as st
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
from googleapiclient.discovery import build
import json
import os
# --- Imports Adicionais para Geração de Arquivos ---
import io
from docx import Document
import re
import urllib3
# --- CONFIGURAÇÃO E ESTILO ---
st.set_page_config(page_title="Cognos Job AI Pro", page_icon="⚡", layout="wide")

st.markdown("""
<style>
    .job-card { background-color: #1e2229; padding: 20px; border-radius: 12px; border: 1px solid #30363d; margin-bottom: 15px; }
    .stButton>button { border-radius: 8px; font-weight: bold; }
    .match-score { font-size: 2em; font-weight: bold; text-align: center; }
    .st-emotion-cache-16txtl3 { padding-top: 2rem; }
</style>
""", unsafe_allow_html=True)
KEYS_FILE = "user_keys.json"

def load_keys_from_file():
    """Carrega as chaves do arquivo JSON se ele existir."""
    if os.path.exists(KEYS_FILE):
        try:
            with open(KEYS_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

# ==============================================================================
# 🚨 CORREÇÃO: MOVA ESTE BLOCO PARA CÁ (ANTES DO SESSION STATE)
# ==============================================================================

# Nome do arquivo onde as chaves serão salvas
KEYS_FILE = "user_keys.json"

def load_keys_from_file():
    """Carrega as chaves do arquivo JSON se ele existir."""
    if os.path.exists(KEYS_FILE):
        try:
            with open(KEYS_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_keys_to_file(g_key, g_cx, gem_key):
    """Salva as chaves no arquivo JSON."""
    data = {"g_key": g_key, "g_cx": g_cx, "gem_key": gem_key}
    with open(KEYS_FILE, "w") as f:
        json.dump(data, f)

# ==============================================================================

# --- INICIALIZAÇÃO DO ESTADO DA SESSÃO ---
# AGORA VAI FUNCIONAR POIS A FUNÇÃO JÁ FOI LIDA ACIMA
saved_keys = load_keys_from_file() 

keys_to_initialize = [
    'g_key', 'g_cx', 'gem_key', 'user_cv', 'search_results', 
    'selected_job', 'job_description', 'analysis_result',
    'cv_text_out', 'cl_text_out', 'inst_out'
]

for key in keys_to_initialize:
    if key not in st.session_state:
        # Se a chave for uma das credenciais e existir no arquivo salvo, usa ela
        if key in saved_keys and saved_keys[key]:
            st.session_state[key] = saved_keys[key]
        else:
            st.session_state[key] = None

# --- FUNÇÕES DE NÚCLEO (COM CACHE) ---
@st.cache_data(show_spinner="Buscando vagas no Google...")
def util_google_search(query, api_key, cx_id):
    """Executa a busca via Google Custom Search API."""
    try:
        service = build("customsearch", "v1", developerKey=api_key)
        res = service.cse().list(q=query, cx=cx_id, num=10, sort='date').execute()
        return res.get('items', [])
    except Exception as e:
        st.error(f"Erro na API do Google: {e}")
        return []



@st.cache_data(show_spinner="Lendo conteúdo da vaga com IA Reader...")

def clean_html_noise(soup):
    """
    Remove poluição visual com lógica baseada em CONTEÚDO, não só classes.
    """
    # 1. Remove tags estruturais inúteis para IA
    for element in soup(['script', 'style', 'noscript', 'iframe', 'svg', 'header', 'footer', 'nav', 'aside', 'form', 'button']):
        element.decompose()

    # 2. COOKIE NUKE: Remove elementos que contenham texto de consentimento
    # Isso resolve o problema do seu Print 1
    blacklist_phrases = ['utilizamos cookies', 'sua privacidade', 'aceitar todos', 'política de privacidade', 'configurações de cookies']
    
    # Varre divs, sections e spans
    for tag in soup.find_all(['div', 'section', 'span', 'p', 'aside']):
        # Se o texto for curto (banner) e tiver palavras chave de cookie
        text_content = tag.get_text(" ", strip=True).lower()
        if len(text_content) < 400 and any(phrase in text_content for phrase in blacklist_phrases):
            tag.decompose()

    return soup

@st.cache_data(show_spinner="Acessando site e quebrando proteções...")
def scrape_job_description(url):
    """
    Extração Resiliente: Foca em trazer TEXTO, ignorando formatação se necessário.
    """
    extracted_text = ""
    
    # Headers simulando um Chrome real no Windows
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }

    try:
        # verify=False pula erros de SSL que bloqueiam muitos scrapers
        response = requests.get(url, headers=headers, timeout=20, verify=False)
        
        # Detecção automática de encoding (resolve caracteres estranhos)
        response.encoding = response.apparent_encoding
        
        if response.status_code != 200:
            return f"Erro ao acessar site: Status {response.status_code}"

        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 1. Limpeza Pesada (Remove o banner de cookie)
        soup = clean_html_noise(soup)
        
        # 2. Estratégia de Extração "Rede de Arrasto"
        # Em vez de tentar achar o container perfeito, vamos pegar todo texto relevante
        # Sites como o da Solor usam 'accordions', o texto está lá, só escondido no CSS.
        # O BeautifulSoup VÊ o texto escondido, nós só precisamos pegar.
        
        text_blocks = []
        # Pega parágrafos, itens de lista e divs de texto
        for tag in soup.find_all(['p', 'li', 'h1', 'h2', 'h3', 'div']):
            text = tag.get_text(" ", strip=True)
            # Filtra linhas muito curtas ou vazias
            if len(text) > 30: 
                text_blocks.append(text)

        # Remove duplicatas mantendo a ordem
        seen = set()
        clean_blocks = [x for x in text_blocks if not (x in seen or seen.add(x))]
        
        extracted_text = "\n\n".join(clean_blocks)

    except Exception as e:
        return f"Erro crítico na extração: {str(e)}. Tente copiar e colar manual."

    # Validação Final
    if len(extracted_text) < 100:
        # Fallback para Jina se o requests falhar
        try:
            return requests.get(f"https://r.jina.ai/{url}", timeout=10).text
        except:
            return "⚠️ O site bloqueou o acesso ou o conteúdo é gerado 100% via Javascript complexo. Por favor, copie o texto da vaga e cole na aba 'Texto Manual'."
        
    return extracted_text

@st.cache_data(show_spinner="Web Specter extraindo dados...")
def scrape_job_description(url):
    """
    Extração Híbrida v2.0:
    Tenta Jina (bom para JS) -> Se falhar ou vier sujo -> Usa Requests com headers Black + Limpeza Cirúrgica.
    """
    extracted_text = ""
    
    # --- TENTATIVA 1: Lógica Manual Robustecida (Prioridade para limpeza local) ---
    # Motivo: O Jina às vezes traz o banner de cookie renderizado. Nossa limpeza local é mais segura.
    try:
        # Headers anti-bloqueio (Black Edition)
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1'
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        
        # Se der erro 403/401 (bloqueio), pula para o Jina
        if response.status_code in [403, 401, 503]:
            raise Exception("Bloqueio de WAF detectado")

        response.encoding = response.apparent_encoding # Corrige acentuação
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 1. Limpa o lixo (Cookies, Menus)
        soup = clean_html_noise(soup)
        
        # 2. Encontra o bloco exato da vaga
        content_block = find_job_content(soup)
        
        # 3. Formata o texto final
        lines = []
        for line in content_block.get_text("\n").splitlines():
            clean_line = line.strip()
            # Filtra linhas inúteis que sobraram
            if len(clean_line) > 2 and clean_line.lower() not in ["aceitar", "fechar", "voltar"]:
                lines.append(clean_line)
        
        extracted_text = "\n".join(lines)

    except Exception as e:
        print(f"Método local falhou ou foi bloqueado: {e}. Tentando Jina...")
        extracted_text = "" # Força fallback

    # --- TENTATIVA 2: Fallback para Jina Reader (Se o local falhar) ---
    if not extracted_text or len(extracted_text) < 150:
        try:
            jina_url = f"https://r.jina.ai/{url}"
            # Headers simples para o Jina
            jheaders = {'User-Agent': 'Mozilla/5.0', 'X-Return-Format': 'markdown'}
            r = requests.get(jina_url, headers=jheaders, timeout=20)
            if r.status_code == 200:
                # Mesmo com Jina, tentamos limpar banners comuns
                raw_text = r.text
                if "cookie" not in raw_text[:200].lower():
                    extracted_text = raw_text
        except:
            pass

    # Validação Final
    if len(extracted_text) < 100:
        return "⚠️ Não foi possível extrair o texto automaticamente (Site protegido ou conteúdo 100% JS). Por favor, copie e cole o texto manualmente na aba ao lado."
        
    return extracted_text

def get_gemini_response(prompt, api_key, model_name="gemini-2.5-pro"):
    """Gera respostas usando a API do Gemini."""
    if not api_key:
        st.error("A chave da API do Gemini não foi configurada.")
        return None
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Erro na API do Gemini: {e}")
        return None

# --- FUNÇÕES AUXILIARES PARA GERAÇÃO DE ARQUIVOS ---
def create_docx(content, title):
    """Cria um documento DOCX em memória a partir de um texto."""
    document = Document()
    document.add_heading(title, level=1)
    # Parser simples para markdown
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            document.add_heading(line[2:], level=2)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            p = document.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.startswith('* ') or line.startswith('✅') or line.startswith('⚠️') or line.startswith('🔑'):
            # Remove o emoji/marcador para um bullet point limpo
            clean_line = line.lstrip('✅⚠️🔑* ')
            document.add_paragraph(clean_line, style='List Bullet')
        else:
            document.add_paragraph(line)
    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

def display_download_buttons(content, title, filename_prefix):
    """Exibe botões de download apenas para DOCX e MD (PDF removido)."""
    st.markdown(f"###### 📥 Baixar {title}")
    d_col1, d_col2 = st.columns(2) # Reduzido para 2 colunas
    with d_col1:
        st.download_button(
            label="Baixar em DOCX", 
            data=create_docx(content, title),
            file_name=f"{filename_prefix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{filename_prefix}",
            use_container_width=True
        )
    with d_col2:
        st.download_button(
            label="Baixar em Markdown (MD)", 
            data=content.encode('utf-8'),
            file_name=f"{filename_prefix}.md", 
            mime="text/markdown",
            key=f"md_{filename_prefix}",
            use_container_width=True
        )
# Nome do arquivo onde as chaves serão salvas

def save_keys_to_file(g_key, g_cx, gem_key):
    """Salva as chaves no arquivo JSON."""
    data = {"g_key": g_key, "g_cx": g_cx, "gem_key": gem_key}
    with open(KEYS_FILE, "w") as f:
        json.dump(data, f)


# --- SIDEBAR DE CONFIGURAÇÃO ---
with st.sidebar:
    st.header("⚙️ Configurações")
    with st.expander("🔑 Gerenciar Chaves de API", expanded=True):
        # O value agora pega do session_state, que pode ter vindo do arquivo
        g_key = st.text_input("Google API Key", value=st.session_state.get('g_key') or '', type="password")
        g_cx = st.text_input("Google CX ID", value=st.session_state.get('g_cx') or '', type="password")
        gem_key = st.text_input("Gemini API Key", value=st.session_state.get('gem_key') or '', type="password")

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("💾 Salvar", use_container_width=True):
                # Atualiza a sessão atual
                st.session_state.g_key = g_key
                st.session_state.g_cx = g_cx
                st.session_state.gem_key = gem_key
                
                # --- AQUI: Salva no arquivo físico ---
                save_keys_to_file(g_key, g_cx, gem_key)
                st.success("Chaves salvas permanentemente!")
                
        with col_btn2:
            if st.button("🗑️ Limpar", use_container_width=True):
                for k in ['g_key', 'g_cx', 'gem_key']:
                    st.session_state[k] = ""
                # Remove o arquivo físico se existir
                if os.path.exists(KEYS_FILE):
                    os.remove(KEYS_FILE)
                st.rerun()
                
        st.header("👤 Seu Perfil")
    st.session_state.user_cv = st.text_area(
        "Cole seu currículo aqui (formato texto)", 
        value=st.session_state.get('user_cv', ''),
        height=300,
        help="Quanto mais detalhado seu currículo, melhor será a análise da IA."
    )
    if st.session_state.user_cv:
        st.success("Currículo carregado. Pronto para análise!")
# --- INTERFACE PRINCIPAL ---
st.title("⚡ Cognos Job AI Pro: Seu Co-piloto de Carreira")

tab_busca, tab_analise, tab_preparacao = st.tabs(["🔍 Busca & Carga de Vagas", "📊 Análise de Match", "📝 Preparação da Candidatura"])

# --- ABA 1: BUSCA & CARGA DE VAGAS ---
with tab_busca:
    st.info("Utilize a busca inteligente ou cole o link direto. O sistema agora suporta sites dinâmicos (Gupy, LinkedIn, etc).")

    # --- SEÇÃO 1: BUSCA GOOGLE ---
    with st.expander("🔍 Buscar Vagas (Google Custom Search)", expanded=True):
        with st.form("search_form"):
            col1, col2 = st.columns([3, 1])
            with col1:
                # Adicionei 'jobs' no placeholder para incentivar termos melhores
                cargo = st.text_input("Cargo / Keywords:", placeholder="Ex: Python Developer Pleno Gupy")
                local = st.text_input("Localização:", placeholder="Ex: Brasil (Remoto)")
            with col2:
                st.write("") 
                st.write("") 
                submit_search = st.form_submit_button("🔎 Buscar", use_container_width=True)

        if submit_search:
            g_key_val = st.session_state.get('g_key')
            g_cx_val = st.session_state.get('g_cx')
            
            if not g_key_val or not g_cx_val:
                st.warning("⚠️ Configure as chaves do Google na sidebar.")
            elif not cargo:
                st.warning("⚠️ Digite um cargo.")
            else:
                # Dork Otimizada para evitar agregadores de spam
                dork_query = f'intitle:"{cargo}" "{local if local else ""}" (site:gupy.io OR site:linkedin.com/jobs OR site:glassdoor.com.br OR site:greenhouse.io OR site:lever.co) -inurl:login'
                
                results = util_google_search(dork_query, g_key_val, g_cx_val)
                if results:
                    st.session_state.search_results = results
                    st.success(f"Encontradas {len(results)} vagas em sites de alta relevância!")
                else:
                    st.warning("Nenhum resultado. Tente remover filtros de localização.")

        # Exibição dos Resultados
        if st.session_state.get('search_results'):
            st.markdown("### 🎯 Selecione para Analisar")
            for i, r in enumerate(st.session_state.search_results):
                with st.container():
                    col_info, col_action = st.columns([4, 1])
                    with col_info:
                        st.markdown(f"**[{r.get('title')}]({r.get('link')})**")
                        st.caption(f"{r.get('displayLink')} • {r.get('snippet')[:100]}...")
                    with col_action:
                        # Botão com chave única e callback visual
                        if st.button("Analisar ⚡", key=f"btn_search_{i}", use_container_width=True):
                            with st.spinner("Lendo vaga (renderizando JS)..."):
                                content = scrape_job_description(r['link'])
                                
                                st.session_state.selected_job = r
                                st.session_state.job_description = content
                                # Reset de estados
                                st.session_state.analysis_result = None
                                st.session_state.cv_text_out = None
                                st.session_state.cl_text_out = None
                                st.session_state.inst_out = None
                                
                                st.toast("Vaga carregada e lida com sucesso!", icon="🚀")
                                st.rerun()

    st.markdown("---")

    # --- SEÇÃO 2: LINK OU TEXTO MANUAL ---
    manual_tab1, manual_tab2 = st.tabs(["🔗 Link Direto (Melhorado)", "📝 Texto Manual"])
    
    with manual_tab1:
        st.write("Cola o link de sites como **Gupy, LinkedIn, Vagas.com**. O sistema tentará ler o conteúdo completo.")
        url_input = st.text_input("URL da Vaga:", placeholder="https://...")
        
        if st.button("🚀 Extrair Conteúdo do Link", use_container_width=True):
            if url_input:
                with st.spinner("Acessando site e extraindo texto..."):
                    job_text = scrape_job_description(url_input)
                    
                    if len(job_text) > 200:
                        st.session_state.selected_job = {
                            'title': f"Vaga Importada: {url_input[:30]}...",
                            'link': url_input,
                            'displayLink': 'Importação Direta',
                            'snippet': 'N/A'
                        }
                        st.session_state.job_description = job_text
                        
                        # Limpa resultados anteriores
                        st.session_state.analysis_result = None
                        st.session_state.cv_text_out = None
                        st.session_state.cl_text_out = None
                        st.session_state.inst_out = None

                        st.success("Conteúdo extraído com sucesso!")
                        with st.expander("Ver conteúdo extraído"):
                            st.markdown(job_text[:1000] + " [...]")
                    else:
                        st.error("Não foi possível ler o texto. O site pode ter bloqueio forte. Tente copiar e colar na aba 'Texto Manual'.")
            else:
                st.warning("Insira uma URL válida.")

    with manual_tab2:
        # (Mantém o código original desta aba, pois já funciona bem para Ctrl+C/Ctrl+V)
        manual_title = st.text_input("Título da Vaga:", key="manual_title_input")
        manual_desc = st.text_area("Descrição Completa:", height=200, key="manual_desc_input")
        if st.button("Salvar Texto Manual", use_container_width=True):
            if manual_title and manual_desc:
                st.session_state.selected_job = {'title': manual_title, 'link': '#', 'displayLink': 'Manual', 'snippet': manual_desc[:50]}
                st.session_state.job_description = manual_desc
                st.session_state.analysis_result = None
                st.toast("Vaga manual salva!", icon="💾")
            else:
                st.warning("Preencha título e descrição.")

# --- ABA 2: ANÁLISE DE MATCH ---
with tab_analise:
    if not st.session_state.get('selected_job'):
        st.info("Selecione ou carregue uma vaga na aba 'Busca & Carga de Vagas' para iniciar a análise.")
    elif not st.session_state.get('user_cv'):
        st.warning("Por favor, cole seu currículo na barra lateral para fazer a análise de match.")
    else:
        st.subheader(f"Análise de Compatibilidade para: {st.session_state.selected_job['title']}")

        if st.button("🤖 Analisar Compatibilidade com IA", use_container_width=True):
            prompt = f"""
            **Tarefa:** Aja como um especialista em recrutamento e seleção (Tech Recruiter). Analise o currículo do candidato e a descrição da vaga fornecida para determinar a compatibilidade.

            **Currículo do Candidato:**
            ---
            {st.session_state.user_cv}
            ---

            **Descrição da Vaga:**
            ---
            Título: {st.session_state.selected_job['title']}
            Conteúdo: {st.session_state.job_description}
            ---

            **Sua Resposta (use estritamente este formato Markdown):**
            1.  **Pontuação de Match:** Forneça uma porcentagem de compatibilidade estimada (de 0% a 100%) e coloque-a em negrito. Ex: **85%**.
            2.  **Resumo da Análise:** Um parágrafo conciso (2-3 linhas) explicando o porquê da pontuação, destacando o alinhamento geral.
            3.  **✅ Pontos Fortes (Match Direto):** Liste em tópicos as 3-5 principais habilidades e experiências do currículo que são mais relevantes e se alinham diretamente aos requisitos da vaga.
            4.  **⚠️ Pontos de Melhoria (Gaps):** Liste em tópicos as competências ou experiências importantes exigidas pela vaga que não estão claras ou estão ausentes no currículo. Para cada ponto, sugira brevemente como o candidato poderia abordar isso (ex: 'Destacar projeto X que envolveu tecnologia Y' ou 'Considerar um curso rápido em Z').
            5.  **🔑 Palavras-chave Essenciais:** Liste de 5 a 7 palavras-chave e tecnologias cruciais da vaga que o candidato deve garantir que estejam presentes em seu currículo e carta de apresentação para passar por sistemas de triagem (ATS).
            """
            with st.spinner("IA analisando seu perfil contra a vaga..."):
                st.session_state.analysis_result = get_gemini_response(prompt, st.session_state.get('gem_key'))

        if st.session_state.get('analysis_result'):
            st.markdown("---")
            st.markdown(st.session_state.analysis_result)
            display_download_buttons(st.session_state.analysis_result, "Análise de Match", "analise_match")


# --- ABA 3: PREPARAÇÃO DA CANDIDATURA ---
with tab_preparacao:
    if not st.session_state.get('selected_job'):
        st.info("Selecione uma vaga e faça a análise para preparar seus documentos.")
    elif not st.session_state.get('user_cv'):
        st.warning("Por favor, cole seu currículo na barra lateral para gerar os documentos.")
    else:
        st.subheader(f"Documentos Otimizados para: {st.session_state.selected_job['title']}")

        col_cv, col_cl, col_ent = st.columns(3)

        with col_cv:
            if st.button("📄 Gerar Currículo Otimizado", use_container_width=True):
                # OTIMIZAÇÃO 3: Prompt melhorado para o currículo
                prompt = f"""
                **Tarefa:** Aja como um Career Coach especialista em otimização de currículos para sistemas ATS (Applicant Tracking System).
                Adapte o currículo original fornecido para se alinhar perfeitamente com a descrição da vaga. O objetivo é destacar as experiências e habilidades mais relevantes, usando verbos de ação e resultados quantificáveis.

                **Regras:**
                1.  **Não invente informações.** Apenas reestruture, refine e dê ênfase ao que já existe no currículo original.
                2.  **Incorpore palavras-chave** da descrição da vaga naturalmente no texto.
                3.  Siga a estrutura profissional abaixo.

                **Meu Currículo Original:**
                ---
                {st.session_state.user_cv}
                ---

                **Descrição da Vaga Alvo:**
                ---
                {st.session_state.job_description}
                ---

                **Resultado Esperado (use este formato Markdown):**
                # [Nome do Candidato]
                [Seu E-mail] | [Seu Telefone] | [Link para LinkedIn/Portfólio]

                ## Resumo Profissional
                (Crie um parágrafo de 3-4 linhas que resuma as qualificações mais importantes do candidato PARA ESTA VAGA, começando com o cargo e anos de experiência.)

                ## Experiência Profissional
                **[Cargo Mais Recente]** | [Nome da Empresa] | [Período]
                *   (Adapte os bullet points para usar verbos de ação e focar em conquistas que correspondam aos requisitos da vaga. Ex: 'Otimizei processos em X%, resultando em Y' em vez de 'Responsável por processos'.)
                *   (Adicione mais bullet points relevantes...)

                **[Cargo Anterior]** | [Nome da Empresa] | [Período]
                *   (Faça o mesmo para as experiências anteriores, priorizando a relevância para a vaga.)

                ## Educação
                **[Curso]** | [Instituição] | [Ano de Conclusão]

                ## Habilidades
                **Hard Skills:** (Liste as habilidades técnicas mais relevantes para a vaga, como linguagens, frameworks, ferramentas).
                **Soft Skills:** (Liste habilidades comportamentais mencionadas ou implícitas na vaga).
                """
                with st.spinner("IA otimizando seu currículo..."):
                    st.session_state.cv_text_out = get_gemini_response(prompt, st.session_state.get('gem_key'))

            if st.session_state.get('cv_text_out'):
                st.text_area("Currículo Otimizado:", st.session_state.cv_text_out, height=400, key="cv_out_area")
                display_download_buttons(st.session_state.cv_text_out, "Currículo", "curriculo_otimizado")

        with col_cl:
            if st.button("✉️ Gerar Carta de Apresentação", use_container_width=True):
                # OTIMIZAÇÃO 3: Prompt melhorado para a carta de apresentação
                prompt = f"""
                **Tarefa:** Aja como um redator profissional e escreva uma carta de apresentação convincente e personalizada para a vaga, usando as informações do currículo do candidato.

                **Regras:**
                1.  A carta deve ser concisa, profissional e entusiástica.
                2.  Siga a estrutura clássica de 3 parágrafos.
                3.  **Não deve repetir o currículo**, mas sim destacar 2-3 realizações chave que se conectam diretamente com os problemas ou necessidades descritas na vaga.

                **Meu Currículo:**
                ---
                {st.session_state.user_cv}
                ---

                **Descrição da Vaga Alvo:**
                ---
                {st.session_state.job_description}
                ---

                **Resultado Esperado (use este formato Markdown):**

                Prezados(as) recrutadores(as) da [Nome da Empresa, se possível inferir da vaga],

                (Parágrafo 1: Introdução. Mencione a vaga para a qual está se candidatando e onde a viu. Expresse forte interesse na oportunidade e na empresa, mostrando que fez uma pesquisa mínima.)

                (Parágrafo 2: Corpo. Esta é a parte principal. Conecte sua experiência diretamente a um ou dois requisitos CRÍTICOS da vaga. Use um exemplo de projeto ou uma conquista do seu currículo para ilustrar como você pode agregar valor. Ex: "No meu papel na Empresa X, liderei um projeto que resultou em [resultado quantificável], o que se alinha diretamente à vossa necessidade de [requisito da vaga].")

                (Parágrafo 3: Conclusão. Reafirme seu entusiasmo pela vaga. Mencione como seus valores ou objetivos se alinham com os da empresa. Termine com uma chamada para ação clara, como 'Tenho grande interesse em discutir como minhas habilidades em [habilidade chave] podem beneficiar sua equipe.')

                Atenciosamente,

                [Seu Nome Completo]
                """
                with st.spinner("IA redigindo sua carta de apresentação..."):
                    st.session_state.cl_text_out = get_gemini_response(prompt, st.session_state.get('gem_key'))

            if st.session_state.get('cl_text_out'):
                st.text_area("Carta de Apresentação:", st.session_state.cl_text_out, height=400, key="cl_out_area")
                display_download_buttons(st.session_state.cl_text_out, "Carta de Apresentação", "carta_apresentacao")

        with col_ent:
            if st.button("💡 Gerar Dicas para Entrevista", use_container_width=True):
                prompt = f"""
                Com base na descrição da vaga e no meu currículo, crie um guia de preparação para a entrevista. O guia deve ser prático e direto.

                **Meu Currículo:**
                ---
                {st.session_state.user_cv}
                ---

                **Descrição da Vaga:**
                ---
                {st.session_state.job_description}
                ---

                **Resultado Esperado (use este formato Markdown):**

                ## Guia de Preparação para Entrevista: [Título da Vaga]

                ### 1. Possíveis Perguntas Técnicas
                *   (Crie 3 perguntas técnicas específicas baseadas nos requisitos mais importantes da vaga, como 'Como você lidaria com [problema técnico descrito na vaga]?').
                *   ...
                *   ...

                ### 2. Possíveis Perguntas Comportamentais
                *   (Crie 3 perguntas comportamentais clássicas, mas com um viés para o contexto da vaga. Ex: 'Descreva uma situação em que você teve que aprender uma nova tecnologia rapidamente.')
                *   ...
                *   ...

                ### 3. Respostas Sugeridas (Método STAR)
                **Para a pergunta:** (Escolha uma das perguntas comportamentais acima).
                *   **Situação:** (Descreva um contexto relevante do currículo do candidato).
                *   **Tarefa:** (Qual era o objetivo ou desafio?).
                *   **Ação:** (Quais passos específicos o candidato tomou? Use verbos de ação).
                *   **Resultado:** (Qual foi o resultado positivo? Quantifique, se possível).

                ### 4. Perguntas para você fazer ao Entrevistador
                *   (Sugira 3 perguntas inteligentes que demonstrem interesse e conhecimento. Ex: 'Quais são os maiores desafios que a equipe enfrenta atualmente e como esta posição ajudaria a superá-los?').
                *   ...
                *   ...
                """
                with st.spinner("IA preparando suas dicas para a entrevista..."):
                    st.session_state.inst_out = get_gemini_response(prompt, st.session_state.get('gem_key'))

            if st.session_state.get('inst_out'):
                st.markdown(st.session_state.inst_out)
                st.markdown("---")
                display_download_buttons(st.session_state.inst_out, "Dicas de Entrevista", "dicas_entrevista")

